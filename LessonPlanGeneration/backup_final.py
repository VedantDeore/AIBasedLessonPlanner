import random
from flask import Flask, make_response, render_template, request, redirect, url_for, session, jsonify, send_from_directory
import json
import os
from datetime import datetime, timedelta
import pandas as pd
import numpy as np

from flask import Flask, render_template, request, redirect, url_for, send_file, flash

from PIL import Image
import pytesseract
import json
import os
import re
import pandas as pd
import datetime as dt


from io import BytesIO
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from PIL import Image, ImageDraw, ImageFont
from docx import Document
import requests


app = Flask(__name__)
app.secret_key = 'your_secret_key'  # For session management



UPLOAD_FOLDER = 'text_files'
ALLOWED_EXTENSIONS = {'txt'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER


# Function to read text files
def read_text_files(directory):
    documents = []
    for filename in os.listdir(directory):
        if filename.endswith('.txt'):
            with open(os.path.join(directory, filename), 'r', encoding='utf-8') as file:
                documents.append(file.read())
    return documents

# ///////////////////////////////////////////////////////////

# Configure Tesseract (adjust path as needed)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Ensure directories exist
os.makedirs('uploads', exist_ok=True)
os.makedirs('data', exist_ok=True)

# OCR and Text Processing Functions
def extract_text_from_image(image_path):
    try:
        image = Image.open(image_path)
        return pytesseract.image_to_string(image)
    except Exception as e:
        return f"Error during OCR: {str(e)}"

def clean_text(raw_text):
    # Remove headers, footers, unnecessary lines
    filtered_text = re.sub(r"\[.*?\]|\bPage\s+\d+.*?\b|Vishwakarma.*?Pune.*?|Issue\s+\d+:.*?\d{2}/\d{2}/\d{2}", "", raw_text)
    filtered_text = re.sub(r"\n\s*\n", "\n\n", filtered_text).strip()
    return filtered_text

def extract_topics_and_subtopics(cleaned_text):
    topics = []
    lines = cleaned_text.split("\n")

    current_topic = None
    for line in lines:
        if line.strip():
            if re.match(r"^[A-Za-z\s]+[\d\s]*[:\-]", line):
                if current_topic:
                    topics.append(current_topic)
                current_topic = {"topic": line.strip(), "subtopics": []}
            elif current_topic:
                current_topic["subtopics"].append(line.strip())
    if current_topic:
        topics.append(current_topic)
    return topics

def generate_lesson_plan(start_date, end_date, lecture_days, public_holidays, flat_lecture_topics):
    current_date = start_date
    lecture_no = 1
    schedule = []

    for lecture_topic in flat_lecture_topics:
        while current_date <= end_date:
            if current_date.weekday() in lecture_days and current_date not in public_holidays:
                schedule.append({
                    "Date": current_date.strftime("%Y-%m-%d"),
                    "Day": current_date.strftime("%A"),
                    "Lecture No": lecture_no,
                    "Topic": lecture_topic
                })
                lecture_no += 1
                current_date += dt.timedelta(days=1)
                break
            current_date += dt.timedelta(days=1)

    return pd.DataFrame(schedule)


import re
import json
import datetime as dt

@app.route('/upload_calendar', methods=['GET', 'POST'])
def upload_calendar():
    if request.method == 'POST':
        # Check if a file was uploaded
        if 'calendar_image' not in request.files:
            return "No file uploaded", 400
        
        calendar_image = request.files['calendar_image']
        
        if calendar_image.filename == '':
            return "No selected file", 400
        
        # Save the uploaded image
        image_path = os.path.join('uploads', 'calendar_image.jpg')
        calendar_image.save(image_path)
        
        # Extract text from image
        calendar_text = extract_text_from_image(image_path)
        
        # Clean the extracted text
        cleaned_calendar_text = clean_text(calendar_text)
        
        # Extract start and end semester dates
        start_of_semester = re.search(r"Start of Semester.*?(\d{1,2}/\d{1,2}/\d{2,4})", cleaned_calendar_text)
        end_of_semester = re.search(r"End Semester Examination.*?(\d{1,2}/\d{1,2}/\d{2,4})", cleaned_calendar_text)
        
        # Extract class test and event dates
        class_test_1 = re.search(r"Class Test 1.*?(\d{1,2}/\d{1,2}/\d{2,4})\s*to\s*(\d{1,2}/\d{1,2}/\d{2,4})", cleaned_calendar_text)
        mid_semester_reviews = re.search(r"Mid Semester Reviews.*?(\d{1,2}/\d{1,2}/\d{2,4})\s*to\s*(\d{1,2}/\d{1,2}/\d{2,4})", cleaned_calendar_text)
        class_test_2 = re.search(r"Class Test 2.*?(\d{1,2}/\d{1,2}/\d{2,4})\s*to\s*(\d{1,2}/\d{1,2}/\d{2,4})", cleaned_calendar_text)
        remedial_teaching = re.search(r"End Semester Remedial Teaching.*?(\d{1,2}/\d{1,2}/\d{2,4})\s*to\s*(\d{1,2}/\d{1,2}/\d{2,4})", cleaned_calendar_text)
        
        # Extract holidays
        holidays = re.findall(r"([A-Za-z\s]+[-/A-Za-z\s]*\d{1,2}/\d{1,2}/\d{2,4})", cleaned_calendar_text)
        
        # Extract event dates (if any)
        event_dates = {
            'start_of_semester': start_of_semester.group(1) if start_of_semester else None,
            'end_of_semester': end_of_semester.group(1) if end_of_semester else None,
            'class_test_1': {'start': class_test_1.group(1), 'end': class_test_1.group(2)} if class_test_1 else None,
            'mid_semester_reviews': {'start': mid_semester_reviews.group(1), 'end': mid_semester_reviews.group(2)} if mid_semester_reviews else None,
            'class_test_2': {'start': class_test_2.group(1), 'end': class_test_2.group(2)} if class_test_2 else None,
            'remedial_teaching': {'start': remedial_teaching.group(1), 'end': remedial_teaching.group(2)} if remedial_teaching else None
        }
        
        # Extract holiday dates
        holiday_dates = []
        for holiday in holidays:
            date_match = re.search(r"(\d{1,2}/\d{1,2}/\d{2,4})", holiday)
            if date_match:
                holiday_dates.append(date_match.group(1))
        
        # Save extracted data to JSON
        with open('data/calendar_data.json', 'w') as f:
            json.dump({
                'calendar_text': cleaned_calendar_text,  # Save the cleaned text
                'raw_text': calendar_text,  # Save the raw OCR text
                'events': event_dates,  # Save the extracted event dates
                'holidays': holiday_dates  # Save the holidays
            }, f, indent=4)
        
        return redirect(url_for('upload_syllabus'))
    
    return render_template('upload_calendar.html')

# Helper function to clean extracted text
def clean_text(raw_text):
    # Remove unwanted characters or fix OCR issues (e.g., incorrect spaces, misinterpreted characters)
    cleaned_text = raw_text.replace("\n", " ").replace("\u00a5", "Y").replace("\u2014", "-")
    
    # Fix spaces, punctuation, and formatting
    cleaned_text = re.sub(r'\s+', ' ', cleaned_text)  # Replace multiple spaces with single space
    cleaned_text = re.sub(r'(\d{1,2}/\d{1,2}/\d{2,4})', r'\1', cleaned_text)  # Ensure date format is consistent
    
    return cleaned_text


# @app.route('/upload_calendar', methods=['GET', 'POST'])
# def upload_calendar():
#     if request.method == 'POST':
#         # Check if a file was uploaded
#         if 'calendar_image' not in request.files:
#             return "No file uploaded", 400
        
#         calendar_image = request.files['calendar_image']
        
#         if calendar_image.filename == '':
#             return "No selected file", 400
        
#         # Save the uploaded image
#         image_path = os.path.join('uploads', 'calendar_image.jpg')
#         calendar_image.save(image_path)
        
#         # Extract text from image
#         calendar_text = extract_text_from_image(image_path)
        
#         # Clean the extracted text
#         cleaned_calendar_text = clean_text(calendar_text)
        
#         # Extract public holidays (if needed)
#         public_holidays = re.findall(r"\b\d{2}/\d{2}/\d{4}\b", cleaned_calendar_text)
#         public_holidays = [dt.datetime.strptime(date, "%d/%m/%Y").date() for date in public_holidays]
        
#         # Save extracted data to JSON
#         with open('data/calendar_data.json', 'w') as f:
#             json.dump({
#                 'calendar_text': cleaned_calendar_text,  # Save the cleaned text
#                 'raw_text': calendar_text,  # Save the raw OCR text
#                 'public_holidays': [str(holiday) for holiday in public_holidays]
#             }, f, indent=4)
        
#         return redirect(url_for('upload_syllabus'))
    
#     return render_template('upload_calendar.html')

@app.route('/upload_syllabus', methods=['GET', 'POST'])
def upload_syllabus():
    if request.method == 'POST':
        # Check if a file was uploaded
        if 'syllabus_image' not in request.files:
            return "No file uploaded", 400
        
        syllabus_image = request.files['syllabus_image']
        
        if syllabus_image.filename == '':
            return "No selected file", 400
        
        # Save the uploaded image
        image_path = os.path.join('uploads', 'syllabus_image.jpg')
        syllabus_image.save(image_path)
        
        # Extract text from image
        syllabus_text = extract_text_from_image(image_path)
        cleaned_syllabus_text = clean_text(syllabus_text)
        
        # Extract topics
        topics = extract_topics_and_subtopics(cleaned_syllabus_text)
        
        # Save extracted data to JSON
        with open('data/syllabus_data.json', 'w') as f:
            json.dump({
                'syllabus_text': cleaned_syllabus_text,
                'topics': topics
            }, f, indent=4)
        
        return redirect(url_for('edit_calendar'))
    
    return render_template('upload_syllabus.html')


@app.route('/edit_calendar', methods=['GET', 'POST'])
def edit_calendar():
    calendar_file = 'data/calendar_data.json'
    
    # Load existing data from JSON
    with open(calendar_file, 'r') as f:
        calendar_data = json.load(f)

    if request.method == 'POST':
        # Update events
        events = {
            "start_of_semester": request.form.get("start_of_semester"),
            "end_of_semester": request.form.get("end_of_semester"),
            "class_test_1": request.form.get("class_test_1"),
            "mid_semester_reviews": request.form.get("mid_semester_reviews"),
            "class_test_2": request.form.get("class_test_2"),
            "remedial_teaching": request.form.get("remedial_teaching"),
        }

        # Update holidays
        holidays = request.form.getlist("holidays[]")
        
        # Update lecture days and times
        lecture_days = request.form.getlist("days[]")
        lecture_times = request.form.getlist("times[]")

        # Save updated data back to JSON
        calendar_data['events'] = events
        calendar_data['holidays'] = holidays
        calendar_data['lecture_days'] = lecture_days
        calendar_data['lecture_times'] = lecture_times

        with open(calendar_file, 'w') as f:
            json.dump(calendar_data, f, indent=4)

        return redirect(url_for('final_plan'))


    # Pre-fill form with existing data
    return render_template('edit_calendar.html', calendar_data=calendar_data, zip=zip)


# /from transformers import pipeline

# Helper function to calculate valid lecture dates
def generate_valid_dates(start_date, end_date, lecture_days, holidays, events):
    start = datetime.strptime(start_date, "%Y-%m-%d")
    end = datetime.strptime(end_date, "%Y-%m-%d")
    day_mapping = {day: i for i, day in enumerate(["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"])}
    lecture_days_indices = [day_mapping[day] for day in lecture_days]
    holidays = [datetime.strptime(holiday, "%d/%m/%y") for holiday in holidays]
    event_dates = [datetime.strptime(date, "%Y-%m-%d") for date in events.values()]
    excluded_dates = set(holidays + event_dates)

    valid_dates = []
    current = start
    while current <= end:
        if current.weekday() in lecture_days_indices and current not in excluded_dates:
            valid_dates.append(current)
        current += timedelta(days=1)
    return valid_dates

# Function to distribute topics across valid lecture dates
def distribute_topics_from_syllabus(syllabus_text, valid_dates):
    """
    Distribute individual topics across valid lecture dates from the syllabus text.
    Each topic is assigned to a single lecture, with larger topics split into smaller parts if necessary.
    """
    lesson_plan = []
    
    # Step 1: Split the syllabus text into individual topics
    # Split syllabus into topics by splitting on the periods.
    topics = [topic.strip() for topic in syllabus_text.split('.') if topic.strip()]
    
    # Step 2: Further split topics with commas into subtopics
    subtopics = []
    for topic in topics:
        # Split the topic by commas if it's a multi-part topic
        parts = [part.strip() for part in topic.split(',') if part.strip()]
        subtopics.extend(parts)
    
    num_dates = len(valid_dates)
    num_subtopics = len(subtopics)

    # Ensure each lecture gets one subtopic
    for i, date in enumerate(valid_dates):
        assigned_subtopic = subtopics[i % num_subtopics]  # Recycle subtopics if there are more dates than subtopics
        lesson_plan.append({
            "date": date.strftime("%Y-%m-%d"),
            "topic": assigned_subtopic
        })

    return lesson_plan



@app.route('/final_plan', methods=['GET', 'POST'])
def final_plan():
    # Load calendar and syllabus data
    with open('data/calendar_data.json', 'r') as f:
        calendar_data = json.load(f)
    with open('data/syllabus_data.json', 'r') as f:
        syllabus_data = json.load(f)

    # Extract data
    start_date = calendar_data['events']['start_of_semester']
    end_date = calendar_data['events']['end_of_semester']
    lecture_days = calendar_data['lecture_days']
    holidays = calendar_data['holidays']
    events = calendar_data['events']
    syllabus_text = syllabus_data.get('syllabus_text', "")

    # Generate valid lecture dates
    valid_dates = generate_valid_dates(start_date, end_date, lecture_days, holidays, events)

    # Initialize lesson_plan
    lesson_plan = []

    # Distribute topics from the syllabus to the valid dates
    lesson_plan = distribute_topics_from_syllabus(syllabus_text, valid_dates)

    # Add additional columns: Method, Student Activity, Assessment Tool, and Remarks
    method_default = "PPT/White Board"
    student_activity_default = "Discussion, Solving of Problem"
    assessment_tools = ["Quiz", "HA", "Tutorial", "ESE"]
    
    # Assuming a generic time slot for now
    time_slot = "10:00 AM - 12:00 PM"
    
    day_mapping = {0: 'Monday', 1: 'Tuesday', 2: 'Wednesday', 3: 'Thursday', 4: 'Friday', 5: 'Saturday', 6: 'Sunday'}

    section_flag = {"I": False, "II": False}  # Track if sections are assigned
    
    for i, entry in enumerate(lesson_plan):
        date_obj = datetime.strptime(entry['date'], "%Y-%m-%d")
        entry['day'] = day_mapping[date_obj.weekday()]
        entry['lecture_number'] = i + 1
        entry['time_slot'] = time_slot
        entry['method'] = method_default
        entry['student_activity'] = student_activity_default
        entry['assessment_tool'] = random.choice(assessment_tools)  # Randomly pick an assessment tool
        entry['remarks'] = ""  # Leave remarks empty initially
        
        # Set Section I and Section II based on the topic
        if "SECTION-I" in entry['topic'] and not section_flag["I"]:
            entry['section'] = "SECTION-I"
            section_flag["I"] = True
        elif "SECTION-II" in entry['topic'] and not section_flag["II"]:
            entry['section'] = "SECTION-II"
            section_flag["II"] = True
        else:
            entry['section'] = ""

    # If it's a POST request, update the lesson plan with the submitted data
    if request.method == 'POST':
        # Handling add or delete actions for rows and columns
        action = request.form.get('action')
        
        if action == 'add_row':
            # Add a row after a specific lecture number
            lecture_number = int(request.form.get('lecture_number'))
            new_row = {
                'date': request.form.get(f'date_{lecture_number + 1}', ''),
                'time_slot': request.form.get(f'time_slot_{lecture_number + 1}', ''),
                'topic': request.form.get(f'topic_{lecture_number + 1}', ''),
                'method': request.form.get(f'method_{lecture_number + 1}', ''),
                'student_activity': request.form.get(f'student_activity_{lecture_number + 1}', ''),
                'assessment_tool': request.form.get(f'assessment_tool_{lecture_number + 1}', ''),
                'remarks': request.form.get(f'remarks_{lecture_number + 1}', ''),
            }
            lesson_plan.insert(lecture_number, new_row)  # Insert row at the specified position
        
        elif action == 'delete_row':
            # Delete a row by lecture number
            lecture_number = int(request.form.get('lecture_number'))
            lesson_plan.pop(lecture_number - 1)  # Delete the row based on lecture number
        
        elif action == 'add_column':
            # Add a column before or after a specific column
            column_name = request.form.get('column_name')
            position = request.form.get('position')  # 'before' or 'after'
            
            for entry in lesson_plan:
                if position == 'before':
                    entry[column_name] = ""  # Add column before
                elif position == 'after':
                    entry[column_name] = ""  # Add column after
        
        elif action == 'delete_column':
            # Delete a specific column from all rows
            column_name = request.form.get('column_name')
            for entry in lesson_plan:
                if column_name in entry:
                    del entry[column_name]  # Delete the column

        # Save updated lesson plan to JSON
        with open('data/lesson_plan.json', 'w') as f:
            json.dump(lesson_plan, f, indent=4)

    # Render the final plan with editable fields
    return render_template('final.html', lesson_plan=lesson_plan)



# /////////   With Quiz for each Sections  ////////////////
# @app.route('/final_plan', methods=['GET', 'POST'])
# def final_plan():
#     # Load calendar and syllabus data
#     with open('data/calendar_data.json', 'r') as f:
#         calendar_data = json.load(f)
#     with open('data/syllabus_data.json', 'r') as f:
#         syllabus_data = json.load(f)

#     # Extract data
#     start_date = calendar_data['events']['start_of_semester']
#     end_date = calendar_data['events']['end_of_semester']
#     lecture_days = calendar_data['lecture_days']
#     holidays = calendar_data['holidays']
#     events = calendar_data['events']
#     syllabus_text = syllabus_data.get('syllabus_text', "")

#     # Generate valid lecture dates
#     valid_dates = generate_valid_dates(start_date, end_date, lecture_days, holidays, events)

#     # Initialize lesson_plan
#     lesson_plan = []

#     # Distribute topics from the syllabus to the valid dates
#     lesson_plan = distribute_topics_from_syllabus(syllabus_text, valid_dates)

#     # Add additional columns: Method, Student Activity, Assessment Tool, and Remarks
#     method_default = "PPT/White Board"
#     student_activity_default = "Discussion, Solving of Problem"
#     assessment_tools = ["Quiz", "HA", "Tutorial", "ESE"]
    
#     # Assuming a generic time slot for now
#     time_slot = "10:00 AM - 12:00 PM"
    
#     day_mapping = {0: 'Monday', 1: 'Tuesday', 2: 'Wednesday', 3: 'Thursday', 4: 'Friday', 5: 'Saturday', 6: 'Sunday'}

#     section_flag = {"I": False, "II": False}  # Track if sections are assigned
    
#     for i, entry in enumerate(lesson_plan):
#         date_obj = datetime.strptime(entry['date'], "%Y-%m-%d")
#         entry['day'] = day_mapping[date_obj.weekday()]
#         entry['lecture_number'] = i + 1
#         entry['time_slot'] = time_slot
#         entry['method'] = method_default
#         entry['student_activity'] = student_activity_default
#         entry['assessment_tool'] = random.choice(assessment_tools)  # Randomly pick an assessment tool
#         entry['remarks'] = ""  # Leave remarks empty initially
        
#         # Set Section I and Section II based on the topic
#         if "SECTION-I" in entry['topic'] and not section_flag["I"]:
#             entry['section'] = "SECTION-I"
#             section_flag["I"] = True
#         elif "SECTION-II" in entry['topic'] and not section_flag["II"]:
#             entry['section'] = "SECTION-II"
#             section_flag["II"] = True
#         else:
#             entry['section'] = ""

#     # If it's a POST request, update the lesson plan with the submitted data
#     if request.method == 'POST':
#         for i, entry in enumerate(lesson_plan):
#             entry['date'] = request.form.get(f'date_{i+1}')
#             entry['time_slot'] = request.form.get(f'time_slot_{i+1}')
#             entry['topic'] = request.form.get(f'topic_{i+1}')
#             entry['method'] = request.form.get(f'method_{i+1}')
#             entry['student_activity'] = request.form.get(f'student_activity_{i+1}')
#             entry['assessment_tool'] = request.form.get(f'assessment_tool_{i+1}')
#             entry['remarks'] = request.form.get(f'remarks_{i+1}')
        
#         # Save updated lesson plan to JSON
#         with open('data/lesson_plan.json', 'w') as f:
#             json.dump(lesson_plan, f, indent=4)

#     # Render the final plan with editable fields
#     return render_template('final.html', lesson_plan=lesson_plan)




# /////////   With Quiz for each row ////////////////


# @app.route('/final_plan', methods=['GET', 'POST'])
# def final_plan():
#     # Load calendar and syllabus data
#     with open('data/calendar_data.json', 'r') as f:
#         calendar_data = json.load(f)
#     with open('data/syllabus_data.json', 'r') as f:
#         syllabus_data = json.load(f)

#     # Extract data
#     start_date = calendar_data['events']['start_of_semester']
#     end_date = calendar_data['events']['end_of_semester']
#     lecture_days = calendar_data['lecture_days']
#     holidays = calendar_data['holidays']
#     events = calendar_data['events']
#     syllabus_text = syllabus_data.get('syllabus_text', "")

#     # Generate valid lecture dates
#     valid_dates = generate_valid_dates(start_date, end_date, lecture_days, holidays, events)

#     # Initialize lesson_plan
#     lesson_plan = []

#     # Distribute topics from the syllabus to the valid dates
#     lesson_plan = distribute_topics_from_syllabus(syllabus_text, valid_dates)

#     # Add additional columns: Method, Student Activity, Assessment Tool, and Remarks
#     method_default = "PPT/White Board"
#     student_activity_default = "Discussion, Solving of Problem"
#     assessment_tools = ["Quiz", "HA", "Tutorial", "ESE"]
    
#     # Assuming a generic time slot for now
#     time_slot = "10:00 AM - 12:00 PM"
    
#     day_mapping = {0: 'Monday', 1: 'Tuesday', 2: 'Wednesday', 3: 'Thursday', 4: 'Friday', 5: 'Saturday', 6: 'Sunday'}

#     for i, entry in enumerate(lesson_plan):
#         date_obj = datetime.strptime(entry['date'], "%Y-%m-%d")
#         entry['day'] = day_mapping[date_obj.weekday()]
#         entry['lecture_number'] = i + 1
#         entry['time_slot'] = time_slot
#         entry['method'] = method_default
#         entry['student_activity'] = student_activity_default
#         entry['assessment_tool'] = random.choice(assessment_tools)  # Randomly pick an assessment tool
#         entry['remarks'] = ""  # Leave remarks empty initially

#     # If it's a POST request, update the lesson plan with the submitted data
#     if request.method == 'POST':
#         for i, entry in enumerate(lesson_plan):
#             entry['date'] = request.form.get(f'date_{i+1}')
#             entry['time_slot'] = request.form.get(f'time_slot_{i+1}')
#             entry['topic'] = request.form.get(f'topic_{i+1}')
#             entry['method'] = request.form.get(f'method_{i+1}')
#             entry['student_activity'] = request.form.get(f'student_activity_{i+1}')
#             entry['assessment_tool'] = request.form.get(f'assessment_tool_{i+1}')
#             entry['remarks'] = request.form.get(f'remarks_{i+1}')
        
#         # Save updated lesson plan to JSON
#         with open('data/lesson_plan.json', 'w') as f:
#             json.dump(lesson_plan, f, indent=4)

#     # Render the final plan with editable fields
#     return render_template('final.html', lesson_plan=lesson_plan)
# /////////////////////////////

# @app.route('/final_plan', methods=['GET', 'POST'])
# def final_plan():
#     # Load calendar and syllabus data
#     with open('data/calendar_data.json', 'r') as f:
#         calendar_data = json.load(f)
#     with open('data/syllabus_data.json', 'r') as f:
#         syllabus_data = json.load(f)

#     # Extract data
#     start_date = calendar_data['events']['start_of_semester']
#     end_date = calendar_data['events']['end_of_semester']
#     lecture_days = calendar_data['lecture_days']
#     holidays = calendar_data['holidays']
#     events = calendar_data['events']
#     syllabus_text = syllabus_data.get('syllabus_text', "")

#     # Generate valid lecture dates
#     valid_dates = generate_valid_dates(start_date, end_date, lecture_days, holidays, events)

#     # Initialize lesson_plan
#     lesson_plan = []

#     # Distribute topics from the syllabus to the valid dates
#     lesson_plan = distribute_topics_from_syllabus(syllabus_text, valid_dates)

#     # Add Day, Lecture Number, and Time Slot to the lesson plan
#     day_mapping = {0: 'Monday', 1: 'Tuesday', 2: 'Wednesday', 3: 'Thursday', 4: 'Friday', 5: 'Saturday', 6: 'Sunday'}
    
#     # Assuming a generic time slot for now
#     time_slot = "10:00 AM - 12:00 PM"
    
#     for i, entry in enumerate(lesson_plan):
#         date_obj = datetime.strptime(entry['date'], "%Y-%m-%d")
#         entry['day'] = day_mapping[date_obj.weekday()]
#         entry['lecture_number'] = i + 1
#         entry['time_slot'] = time_slot

#     # If it's a POST request, update the lesson plan with the submitted data
#     if request.method == 'POST':
#         for i, entry in enumerate(lesson_plan):
#             entry['date'] = request.form.get(f'date_{i+1}')
#             entry['time_slot'] = request.form.get(f'time_slot_{i+1}')
#             entry['topic'] = request.form.get(f'topic_{i+1}')
        
#         # Save updated lesson plan to JSON
#         with open('data/lesson_plan.json', 'w') as f:
#             json.dump(lesson_plan, f, indent=4)

#     # Render the final plan with editable fields
#     return render_template('final.html', lesson_plan=lesson_plan)



# @app.route('/final_plan')
# def final_plan():
#     # Load calendar and syllabus data
#     with open('data/calendar_data.json', 'r') as f:
#         calendar_data = json.load(f)
#     with open('data/syllabus_data.json', 'r') as f:
#         syllabus_data = json.load(f)

#     # Extract data
#     start_date = calendar_data['events']['start_of_semester']
#     end_date = calendar_data['events']['end_of_semester']
#     lecture_days = calendar_data['lecture_days']
#     holidays = calendar_data['holidays']
#     events = calendar_data['events']
#     syllabus_text = syllabus_data.get('syllabus_text', "")

#     # Generate valid lecture dates
#     valid_dates = generate_valid_dates(start_date, end_date, lecture_days, holidays, events)

#     # Initialize lesson_plan
#     lesson_plan = []

#     # Attempt Gemini Chatbot API call
#     query = "Create a structured lesson plan with topics distributed evenly across valid lecture days."
#     api_payload = {
#         'start_date': start_date,
#         'end_date': end_date,
#         'lecture_days': lecture_days,
#         'holidays': holidays,
#         'events': events,
#         'syllabus_text': syllabus_text,
#         'query': query
#     }

#     try:
#         response = requests.post('https://gemini-chatbot-api-url.com/generate_plan', json=api_payload)
#         if response.status_code == 200:
#             api_response = response.json()
#             lesson_plan = api_response.get('lesson_plan', [])
#         else:
#             print(f"API error: {response.status_code} - {response.text}")
#     except Exception as e:
#         print(f"API call failed: {e}")

#     # Fallback to local logic if API fails or returns no data
#     if not lesson_plan:
#         lesson_plan = distribute_topics_from_syllabus(syllabus_text, valid_dates)

#     # Ensure lesson_plan is valid before saving
#     if not lesson_plan:
#         lesson_plan = [{"date": "N/A", "topic": "No topics available or scheduling failed"}]

#     # Save lesson plan to JSON
#     with open('data/lesson_plan.json', 'w') as f:
#         json.dump(lesson_plan, f, indent=4)

#     # Render the final plan
#     return render_template('final.html', lesson_plan=lesson_plan)


@app.route('/download_lesson_plan', methods=['GET', 'POST'])
def download_lesson_plan():
    # Read the lesson plan data
    with open('data/lesson_plan.json', 'r') as f:
        lesson_plan = json.load(f)
    
    # Handle download based on format chosen
    if request.method == 'POST':
        format_choice = request.form['format_choice']
        
        if format_choice == 'json':
            return download_json(lesson_plan)
        elif format_choice == 'pdf':
            return download_pdf(lesson_plan)
        elif format_choice == 'xlsx':
            return download_xlsx(lesson_plan)
        elif format_choice == 'img':
            return download_image(lesson_plan)
        elif format_choice == 'word':
            return download_word(lesson_plan)
    
    return render_template('download.html')

def download_json(lesson_plan):
    response = make_response(json.dumps(lesson_plan, indent=4))
    response.headers['Content-Type'] = 'application/json'
    response.headers['Content-Disposition'] = 'attachment; filename=lesson_plan.json'
    return response

def download_pdf(lesson_plan):
    # Create a PDF document
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    c.setFont("Helvetica", 12)

    # Write lesson plan data to PDF
    y_position = 750
    c.drawString(30, y_position, "Lesson Plan")
    y_position -= 20

    for entry in lesson_plan:
        c.drawString(30, y_position, f"Lecture: {entry['lecture_number']} - Topic: {entry['topic']}")
        y_position -= 20
        if y_position < 40:
            c.showPage()
            y_position = 750
        c.drawString(30, y_position, f"Date: {entry['date']} - Time: {entry['time_slot']}")
        y_position -= 20

    c.save()
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="lesson_plan.pdf", mimetype="application/pdf")
def download_xlsx(lesson_plan):
    # Create a new Excel workbook and sheet
    wb = Workbook()
    ws = wb.active
    ws.title = "Lesson Plan"
    
    # Create headers
    headers = ["Lecture Number", "Date", "Time Slot", "Topic", "Method", "Student Activity", "Assessment Tool", "Remarks"]
    ws.append(headers)
    
    # Add lesson plan data
    for entry in lesson_plan:
        row = [
            entry.get('lecture_number', ''),
            entry.get('date', ''),
            entry.get('time_slot', ''),
            entry.get('topic', ''),
            entry.get('method', ''),  # Default to empty string if key doesn't exist
            entry.get('student_activity', ''),  # Default to empty string if key doesn't exist
            entry.get('assessment_tool', ''),  # Default to empty string if key doesn't exist
            entry.get('remarks', '')  # Default to empty string if key doesn't exist
        ]
        ws.append(row)
    
    # Save to a BytesIO object
    from io import BytesIO
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    return send_file(output, as_attachment=True, download_name="lesson_plan.xlsx", mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

def download_image(lesson_plan):
    # Create a simple image from the lesson plan data
    img = Image.new('RGB', (800, 1200), color='white')
    d = ImageDraw.Draw(img)
    font = ImageFont.load_default()

    y_position = 10
    for entry in lesson_plan:
        d.text((10, y_position), f"Lecture {entry['lecture_number']}: {entry['topic']} - {entry['date']}", font=font, fill=(0, 0, 0))
        y_position += 30

    # Save image to a BytesIO object
    img_io = BytesIO()
    img.save(img_io, 'PNG')
    img_io.seek(0)
    
    return send_file(img_io, as_attachment=True, download_name="lesson_plan.png", mimetype="image/png")

def download_word(lesson_plan):
    # Create a Word document
    doc = Document()
    doc.add_heading('Lesson Plan', 0)
    
    for entry in lesson_plan:
        doc.add_paragraph(f"Lecture {entry['lecture_number']}: {entry['topic']}")
        doc.add_paragraph(f"Date: {entry['date']} - Time Slot: {entry['time_slot']}")
        doc.add_paragraph(f"Method: {entry['method']}")
        doc.add_paragraph(f"Student Activity: {entry['student_activity']}")
        doc.add_paragraph(f"Assessment Tool: {entry['assessment_tool']}")
        doc.add_paragraph(f"Remarks: {entry['remarks']}")
        doc.add_paragraph()
    
    # Save the document to a BytesIO object
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return send_file(doc_io, as_attachment=True, download_name="lesson_plan.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# @app.route('/download_lesson_plan')
# def download_lesson_plan():
#     # Read the lesson plan data
#     with open('data/lesson_plan.json', 'r') as f:
#         lesson_plan = json.load(f)
    
#     # Convert the lesson plan to JSON and return as a downloadable file
#     response = make_response(json.dumps(lesson_plan, indent=4))
#     response.headers['Content-Type'] = 'application/json'
#     response.headers['Content-Disposition'] = 'attachment; filename=lesson_plan.json'
#     return response

# @app.route('/download_lesson_plan')
# def download_lesson_plan():
#     return send_file('data/lesson_plan.json', 
#                      mimetype='application/json', 
#                      as_attachment=True, 
#                      download_name='lesson_plan.json')

# ///////////////////////////////////////////////////////////


# Helper functions
def load_json(filename):
    if os.path.exists(filename):
        with open(filename, 'r') as f:
            return json.load(f)
    return {}

def save_json(filename, data):
    with open(filename, 'w') as f:
        json.dump(data, f, indent=4)
def load_user_credentials():
    if os.path.exists('user_credentials.json'):
        try:
            with open('user_credentials.json', 'r') as f:
                data = f.read()
                if not data.strip():  # If file is empty or only contains whitespace
                    return {}  # Return an empty dictionary
                return json.loads(data)  # Load the JSON content
        except json.JSONDecodeError:
            # Handle the case where the file contains invalid JSON
            print("Error: The user_credentials.json file is invalid.")
            return {}
    return {}  # If file doesn't exist, return an empty dictionary

# Save user credentials
def save_user_credentials(data):
    with open('user_credentials.json', 'w') as f:
        json.dump(data, f, indent=4)



@app.route('/')
def home():
    username = None
    
    if 'user_email' in session:
        username = session['user_email']
        return render_template('index.html', username=username)

    return render_template('index.html')



@app.route('/lessonplanner')
def lessonplanner():
    username = None
    
    if 'user_email' in session:
        username = session['user_email']
        

        return render_template('lessonplanner.html', username=username)

    return render_template('index.html')

@app.route('/profile')
def profile():
    if 'user_email' in session:
        user_email = session['user_email']
        user_credentials = load_user_credentials()

        if user_email in user_credentials:
            user_data = user_credentials[user_email]  # Fetch user data
            user_email = user_email
            return render_template('profile.html', user=user_data, user_email = user_email)

    return redirect(url_for('login'))  # Redirect to login if not logged in

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')

        user_credentials = load_user_credentials()
        if email not in user_credentials or user_credentials[email]['password'] != password:
            return render_template('login.html', error='Invalid email or password!')

        session['user_email'] = email
        session['user_type'] = 'teacher'  # Assuming all registered users are teachers
        return redirect(url_for('home'))

    return render_template('login.html')


# Registration route for teachers
@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        first_name = request.form['first_name']
        last_name = request.form['last_name']
        phone = request.form['phone']
        email = request.form['email']
        password = request.form['password']
        confirm_password = request.form['confirm_password']

        if password != confirm_password:
            return 'Passwords do not match!'

        user_credentials = load_user_credentials()

        if email in user_credentials:
            return 'Email already registered!'

        # Add teacher data
        user_credentials[email] = {
            'first_name': first_name,
            'last_name': last_name,
            'phone': phone,
            'password': password
        }

        save_user_credentials(user_credentials)
        return redirect(url_for('login'))

    return render_template('register.html')


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))


if __name__ == '__main__':
    app.run(debug=True, port=8000)
